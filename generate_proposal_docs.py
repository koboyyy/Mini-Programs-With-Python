from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

# Daftar sponsor
sponsors = [
    "Pemuda Grafika", "Sejarah Baru", "PT BLJ", "D'ulek Resto", "Mandiri Seven",
    "Bos Salad", "Ketum Ketum Terdahulu", "The Malique Kost", "Bank BRI",
    "Raysha Kos", "Mikro Jaya", "Rupat Print", "Warunk Start Up", "Cafe Leccata",
    "Cafe Talacia", "Cafe Floor", "Cafe Zamatra", "Cafe Felicity", "Cafe Qeela DJ",
    "Bank BSI", "Martias Digital Printing", "Kedai Kopi Kembodja", "Sekawan Coffee",
    "Cawan Kopi", "Graphic Photo Studio", "Owner Tara Rias Pengantin"
]

# Jalur ke file template
template_path = "PROPOSAL_SPONSORSHIP_KESENIAN_BATHIN_ALAM.docx"

for sponsor in sponsors:
    # Muat dokumen template
    doc = Document(template_path)
    
    # Cari dan ganti nama penerima di bagian kepala
    for paragraph in doc.paragraphs:
        if "Pemuda Grafika" in paragraph.text:
            # Simpan teks lain dalam paragraf
            original_text = paragraph.text
            new_text = original_text.replace("Pemuda Grafika", sponsor)
            
            # Kosongkan paragraf
            paragraph.clear()
            
            # Tambahkan teks sebelum nama sponsor (jika ada)
            before_text = original_text[:original_text.find("Pemuda Grafika")]
            if before_text:
                run = paragraph.add_run(before_text)
                # Asumsikan font asli dipertahankan untuk teks sebelumnya
                run.font.name = 'Times New Roman'  # Sesuaikan jika font asli berbeda
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(12)
            
            # Tambahkan nama sponsor dengan font Times New Roman ukuran 12
            run = paragraph.add_run(sponsor)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # Untuk mendukung karakter non-ASCII
            run.font.size = Pt(12)
            
            # Tambahkan teks setelah nama sponsor (jika ada)
            after_text = original_text[original_text.find("Pemuda Grafika") + len("Pemuda Grafika"):]
            if after_text:
                run = paragraph.add_run(after_text)
                run.font.name = 'Times New Roman'  # Sesuaikan jika font asli berbeda
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(12)
    
    # Simpan dokumen baru
    file_name = f"Proposal_Sponsorship_{sponsor.replace(' ', '_')}.docx"
    doc.save(file_name)
    print(f"File dibuat: {file_name}")